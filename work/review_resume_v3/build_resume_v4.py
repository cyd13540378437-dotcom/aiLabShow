from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path("/Users/caoyudi/Documents/Codex/2026-07-03/ni-xai")
SRC_DOCX = Path("/Users/caoyudi/Downloads/曹宇迪简历.docx")
OUT_DIR = ROOT / "outputs"
PDF_OUT = OUT_DIR / "曹宇迪简历-浅色视觉版-v4.pdf"
HTML_OUT = OUT_DIR / "曹宇迪简历-浅色视觉版-v4.html"

FONT = "STHeiti"
pdfmetrics.registerFont(TTFont(FONT, "/System/Library/Fonts/STHeiti Medium.ttc"))

PAGE_W, PAGE_H = A4
INK = colors.HexColor("#263845")
TITLE = colors.HexColor("#1F3542")
MUTED = colors.HexColor("#5F6F7B")
FAINT = colors.HexColor("#EFF5F8")
CARD = colors.HexColor("#F8FBFD")
LINE = colors.HexColor("#D8E4EA")
BLUE = colors.HexColor("#2D5668")
BLUE_2 = colors.HexColor("#6F8FA0")
WHITE = colors.white


def clean(text: str) -> str:
    text = " ".join(text.split())
    text = text.replace("负责职责说明：", "负责。")
    text = text.replace(" 半天以内", "半天以内")
    return text


docx = DocxDocument(str(SRC_DOCX))
texts = [clean(p.text.strip()) for p in docx.paragraphs]
headers = [[clean(c.text) for c in t.rows[0].cells] for t in docx.tables]

name = texts[0]
role = "B端 AI 应用产品负责人 / 企业级 AI 产品负责人 / Agent 产品负责人"
contact = "男 丨 33岁 丨 12年工作经验 丨 期望薪资：面议 丨 TEL：13540378437 丨 Email：13540378437@163.com"
summary = texts[4]

core_abilities = [
    ("01 复杂业务抽象建模", "将复杂业务场景抽象为对象、流程、指标、关系、策略与权限边界，统一客户、产品、研发之间的业务语言，降低方案设计与研发实现过程中的沟通成本。"),
    ("02 系统性产品架构能力", "从业务目标、对象模型、主流程、状态流转、权限体系、数据结构、配置策略和异常兜底等维度进行系统设计，避免功能堆叠并提升复用能力。"),
    ("03 AI/Agent应用产品化", "从业务价值、数据基础、模型边界、执行风险和人工审核成本出发，设计 Context、Workflow、Agent 执行链路与可观测机制。"),
    ("04 产品规划与交付管理", "基于客户痛点、业务价值、技术可行性和交付成本制定路线图、版本计划和优先级，推动需求、方案、研发、验收到上线迭代闭环。"),
    ("05 团队管理与跨组织落地", "围绕产品目标、业务目标和团队成长目标建立目标拆解、过程跟踪、评审决策和复盘改进机制，推动方案转化为可落地成果。"),
]

br_header = ["博锐尚格科技股份有限公司", "企业级AI应用与数字化产品负责人", "2021.02 - 至今"]
br_summary = texts[14]
br_duties = texts[15:19]
outcomes = [
    (texts[20], [texts[21]]),
    (texts[22], [texts[23]]),
    (texts[24], [texts[25]]),
    (texts[26], [texts[27]]),
    (texts[28], [texts[29]]),
    (texts[30], [texts[31]]),
    (texts[32], [texts[33], texts[34], texts[35], texts[36], texts[37]]),
    (texts[38], [texts[39]]),
    (texts[40], [texts[41]]),
    (texts[42], [texts[43]]),
    (texts[44], [texts[45]]),
]
ali_header = ["阿里巴巴本地生活集团客如云", "产研中心PaaS产品群 | P7产品经理 / Scrum Master", "2019.08 - 2021.02"]
ali = texts[46:52]
mid_header = ["成都镁克孵化器 / 四川兑商宝 / 住梦网络", "产品经理 / 产品总监", "2016.08 - 2019.08"]
mid = texts[52:55]
grid_header = ["中电启明星（国家电网产业集团）", "产品经理", "2015.08 - 2016.08"]
grid = texts[55:58]
early_header = headers[4]
early = texts[58:60]
edu = headers[5]


base = getSampleStyleSheet()


def style(name, size, leading, color=INK, align=TA_LEFT, before=0, after=0):
    return ParagraphStyle(
        name,
        parent=base["Normal"],
        fontName=FONT,
        fontSize=size,
        leading=leading,
        textColor=color,
        alignment=align,
        spaceBefore=before,
        spaceAfter=after,
    )


S = {
    "name": style("name", 25, 30, TITLE, after=2),
    "role": style("role", 11, 14, TITLE, after=3),
    "contact": style("contact", 8.6, 11, MUTED, after=6),
    "h2": style("h2", 12.2, 15, TITLE, after=5),
    "h3": style("h3", 9.4, 12, BLUE, after=3),
    "body": style("body", 8.35, 11.25, INK, after=1.7),
    "small": style("small", 7.75, 10.4, INK, after=1.2),
    "muted": style("muted", 7.3, 9.5, MUTED, after=0),
    "company": style("company", 9.8, 12, TITLE, after=0),
    "date": style("date", 8.2, 10, MUTED, TA_RIGHT),
    "num": style("num", 8.2, 10, BLUE, TA_CENTER),
    "out_title": style("out_title", 8.45, 10.8, TITLE, after=1.4),
    "out_body": style("out_body", 7.62, 9.9, INK, after=1),
}


def P(text: str, key="body"):
    return Paragraph(escape(text), S[key])


def on_page(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(colors.white)
    canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    canvas.setStrokeColor(BLUE)
    canvas.setLineWidth(2.6)
    if doc.page == 1:
        canvas.line(16 * mm, PAGE_H - 15 * mm, PAGE_W - 16 * mm, PAGE_H - 15 * mm)
    canvas.restoreState()


def section(title: str):
    return Table(
        [[Paragraph(escape(title), S["h2"])]],
        colWidths=[178 * mm],
        style=TableStyle(
            [
                ("LINEBEFORE", (0, 0), (0, 0), 3, BLUE),
                ("LINEBELOW", (0, 0), (-1, -1), 0.5, LINE),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        ),
    )


def callout(text: str):
    return Table(
        [[P(text, "body")]],
        colWidths=[178 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), CARD),
                ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                ("LINEBEFORE", (0, 0), (0, -1), 3, BLUE),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        ),
    )


def ability_grid():
    rows = []
    for i in range(0, 4, 2):
        row = []
        for title_, body in core_abilities[i : i + 2]:
            row.append([Paragraph(f"<b>{escape(title_)}：</b> {escape(body)}", S["small"])])
        rows.append(row)
    rows.append([[Paragraph(f"<b>{escape(core_abilities[4][0])}：</b> {escape(core_abilities[4][1])}", S["small"])], ""])
    t = Table(rows, colWidths=[87 * mm, 87 * mm])
    t.setStyle(
        TableStyle(
            [
                ("SPAN", (0, 2), (1, 2)),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FBFCFD")),
                ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#E4EDF2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return t


def company_row(header):
    return Table(
        [[Paragraph(escape(header[0]), S["company"]), Paragraph(escape(header[1]), S["company"]), Paragraph(escape(header[2]), S["date"])]],
        colWidths=[68 * mm, 74 * mm, 36 * mm],
        style=TableStyle(
            [
                ("LINEBELOW", (0, 0), (-1, -1), 0.8, BLUE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        ),
    )


def bullets(items, width=176 * mm):
    rows = [[Paragraph("•", S["small"]), P(i, "small")] for i in items]
    t = Table(rows, colWidths=[5 * mm, width - 5 * mm])
    t.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    return t


def group_label(text: str):
    return Table(
        [[Paragraph(escape(text), S["h3"])]],
        colWidths=[178 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), FAINT),
                ("LINEBEFORE", (0, 0), (0, -1), 2.5, BLUE_2),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        ),
    )


def outcome(num: int, title: str, bodies: list[str]):
    clean_title = title.split("、", 1)[1] if "、" in title else title
    body = [Paragraph(escape(clean_title), S["out_title"])]
    for b in bodies:
        body.append(Paragraph(escape(b), S["out_body"]))
    t = Table(
        [[Paragraph(str(num), S["num"]), body]],
        colWidths=[9 * mm, 167 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF2F6")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#E2EBF0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (0, 0), 2),
                ("RIGHTPADDING", (0, 0), (0, 0), 2),
                ("TOPPADDING", (0, 0), (0, 0), 6),
                ("LEFTPADDING", (1, 0), (1, 0), 8),
                ("RIGHTPADDING", (1, 0), (1, 0), 4),
                ("TOPPADDING", (1, 0), (1, 0), 5),
                ("BOTTOMPADDING", (1, 0), (1, 0), 5),
            ]
        ),
    )
    return t


def edu_works():
    return Table(
        [
            [
                [Paragraph("教育经历", S["h3"]), P("IIT-伊利诺伊理工大学 | MBA | 在读", "small"), P("成都信息工程学院银杏酒店管理学院 | 信息系统与信息管理 | 本科 | 2011 - 2015", "small")],
                [Paragraph("个人作品", S["h3"]), P("Design Harness | AI 产品设计工程记忆系统 / Agent 工作流产品原型", "small"), P("作品链接：待补充", "small")],
            ]
        ],
        colWidths=[87 * mm, 87 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), CARD),
                ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, LINE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        ),
    )


def build_pdf():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = BaseDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=14 * mm,
        bottomMargin=12 * mm,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="page", frames=[frame], onPage=on_page)])
    story = []

    story.append(Paragraph(name, S["name"]))
    story.append(Paragraph(f"希望岗位：{escape(role)}", S["role"]))
    story.append(Paragraph(escape(contact), S["contact"]))
    story.append(section("职业总结"))
    story.append(callout(summary))
    story.append(Spacer(1, 6))
    story.append(section("核心能力"))
    story.append(ability_grid())
    story.append(Spacer(1, 8))
    story.append(section("工作经历"))
    story.append(company_row(br_header))
    story.append(Paragraph("总结：", S["h3"]))
    story.append(callout(br_summary))
    story.append(Spacer(1, 6))
    story.append(Paragraph("职责范围：", S["h3"]))
    story.append(bullets(br_duties))
    story.append(PageBreak())

    story.append(Spacer(1, 5))
    story.append(Paragraph("关键成果：", S["h3"]))
    story.append(group_label("FMEA 知识工程与 AI 应用冷启动"))
    for idx in [1, 2]:
        story.append(outcome(idx, *outcomes[idx - 1]))
    story.append(group_label("业务配置 / 管理 / 执行闭环与规模化落地"))
    for idx in [3, 4, 6, 7]:
        story.append(outcome(idx, *outcomes[idx - 1]))
    story.append(PageBreak())

    story.append(group_label("AI 复盘与 Agent 调度治理"))
    for idx in [5, 8, 9, 10, 11]:
        story.append(outcome(idx, *outcomes[idx - 1]))
    story.append(Spacer(1, 6))
    story.append(company_row(ali_header))
    story.append(bullets(ali))
    story.append(Spacer(1, 5))
    story.append(company_row(mid_header))
    story.append(bullets(mid))
    story.append(Spacer(1, 5))

    story.append(company_row(grid_header))
    story.append(bullets(grid))
    story.append(Spacer(1, 5))
    story.append(company_row(early_header))
    story.append(bullets(early))
    story.append(Spacer(1, 8))
    story.append(section("教育经历 / 个人作品"))
    story.append(edu_works())

    doc.build(story)


def build_html():
    def ps(items):
        return "".join(f"<li>{escape(i)}</li>" for i in items)

    def out_html(idx):
        t, bs = outcomes[idx - 1]
        t = t.split("、", 1)[1] if "、" in t else t
        return f"<article><i>{idx}</i><div><h4>{escape(t)}</h4>{''.join(f'<p>{escape(b)}</p>' for b in bs)}</div></article>"

    html = f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"/><title>曹宇迪简历-浅色视觉版-v4</title>
<style>
body{{margin:0;background:#edf3f6;color:#263845;font-family:-apple-system,BlinkMacSystemFont,"PingFang SC","Microsoft YaHei",sans-serif}}
.page{{width:210mm;min-height:297mm;margin:18px auto;background:white;box-shadow:0 12px 36px rgba(31,53,66,.12);padding:16mm}}
h1{{font-size:42px;margin:0;color:#1F3542}} h2{{font-size:21px;border-bottom:1px solid #D8E4EA;border-left:5px solid #2D5668;padding:4px 0 7px 10px}} h3{{font-size:15px;color:#2D5668}}
.role{{font-size:19px;font-weight:700;color:#1F3542}} .contact{{color:#5F6F7B;margin:8px 0 18px}}
.callout{{background:#F8FBFD;border:1px solid #D8E4EA;border-left:5px solid #2D5668;padding:14px 16px;line-height:1.75}}
.grid{{display:grid;grid-template-columns:1fr 1fr;gap:8px}} .card{{background:#FBFCFD;border:1px solid #D8E4EA;border-radius:8px;padding:12px;line-height:1.7}} .card:last-child{{grid-column:1/3}}
.company{{display:grid;grid-template-columns:1.25fr 1.15fr auto;border-bottom:2px solid #2D5668;font-weight:700;padding:10px 0 8px;margin-top:12px}}
ul{{padding-left:20px;line-height:1.75}} .group{{background:#EFF5F8;border-left:4px solid #6F8FA0;padding:8px 12px;margin-top:12px;font-weight:700;color:#2D5668}}
article{{display:grid;grid-template-columns:34px 1fr;border-bottom:1px solid #E2EBF0;padding:9px 0}} article i{{background:#EAF2F6;border-radius:20px;text-align:center;padding-top:6px;font-style:normal;color:#2D5668}} article h4{{margin:0 0 6px}} article p{{margin:0;line-height:1.65}}
.dual{{display:grid;grid-template-columns:1fr 1fr;gap:12px}} .box{{background:#F8FBFD;border:1px solid #D8E4EA;border-radius:8px;padding:12px}}
@media print{{body{{background:white}}.page{{margin:0;box-shadow:none;page-break-after:always}}}}
</style></head><body>
<section class="page"><h1>{escape(name)}</h1><div class="role">希望岗位：{escape(role)}</div><div class="contact">{escape(contact)}</div><h2>职业总结</h2><div class="callout">{escape(summary)}</div><h2>核心能力</h2><div class="grid">{''.join(f'<div class="card"><b>{escape(t)}：</b>{escape(b)}</div>' for t,b in core_abilities)}</div><h2>工作经历</h2><div class="company"><span>{escape(br_header[0])}</span><span>{escape(br_header[1])}</span><span>{escape(br_header[2])}</span></div><h3>总结：</h3><div class="callout">{escape(br_summary)}</div><h3>职责范围：</h3><ul>{ps(br_duties)}</ul></section>
<section class="page"><h3>关键成果：</h3><div class="group">FMEA 知识工程与 AI 应用冷启动</div>{out_html(1)}{out_html(2)}<div class="group">业务配置 / 管理 / 执行闭环与规模化落地</div>{out_html(3)}{out_html(4)}{out_html(6)}{out_html(7)}</section>
<section class="page"><div class="group">AI 复盘与 Agent 调度治理</div>{out_html(5)}{out_html(8)}{out_html(9)}{out_html(10)}{out_html(11)}<div class="company"><span>{escape(ali_header[0])}</span><span>{escape(ali_header[1])}</span><span>{escape(ali_header[2])}</span></div><ul>{ps(ali)}</ul><div class="company"><span>{escape(mid_header[0])}</span><span>{escape(mid_header[1])}</span><span>{escape(mid_header[2])}</span></div><ul>{ps(mid)}</ul><div class="company"><span>{escape(grid_header[0])}</span><span>{escape(grid_header[1])}</span><span>{escape(grid_header[2])}</span></div><ul>{ps(grid)}</ul><div class="company"><span>{escape(early_header[0])}</span><span>{escape(early_header[1])}</span><span>{escape(early_header[2])}</span></div><ul>{ps(early)}</ul><h2>教育经历 / 个人作品</h2><div class="dual"><div class="box"><h3>教育经历</h3><p>IIT-伊利诺伊理工大学 | MBA | 在读</p><p>成都信息工程学院银杏酒店管理学院 | 信息系统与信息管理 | 本科 | 2011 - 2015</p></div><div class="box"><h3>个人作品</h3><p>Design Harness | AI 产品设计工程记忆系统 / Agent 工作流产品原型</p><p>作品链接：待补充</p></div></div></section>
</body></html>"""
    HTML_OUT.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_html()
    build_pdf()
    print(HTML_OUT)
    print(PDF_OUT)
