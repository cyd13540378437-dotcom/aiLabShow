from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path("/Users/caoyudi/Documents/Codex/2026-07-03/ni-xai")
SRC = Path("/Users/caoyudi/Downloads/曹宇迪简历.docx")
OUT = ROOT / "outputs"
PDF = OUT / "曹宇迪简历-系统架构视觉版.pdf"
HTML = OUT / "曹宇迪简历-系统架构视觉版.html"

FONT = "STHeiti"
pdfmetrics.registerFont(TTFont(FONT, "/System/Library/Fonts/STHeiti Medium.ttc"))

PAGE_W, PAGE_H = A4

INK = colors.HexColor("#151A1F")
GRAPHITE = colors.HexColor("#25313B")
MUTED = colors.HexColor("#64717D")
HAIR = colors.HexColor("#D9E1E7")
PAPER = colors.HexColor("#FAFBFC")
PANEL = colors.HexColor("#F2F6F8")
TEAL = colors.HexColor("#0F766E")
CYAN = colors.HexColor("#0E7490")
BLUE = colors.HexColor("#1D4E89")
AMBER = colors.HexColor("#A16207")
PLUM = colors.HexColor("#6D4C7D")
WHITE = colors.white


def clean(s: str) -> str:
    s = " ".join(s.split())
    s = s.replace("负责职责说明：", "负责。")
    s = s.replace(" 半天以内", "半天以内")
    return s


def load():
    doc = DocxDocument(str(SRC))
    texts = [clean(p.text.strip()) for p in doc.paragraphs]
    tables = [[clean(c.text) for c in t.rows[0].cells] for t in doc.tables]
    return texts, tables


texts, headers = load()

data = {
    "name": texts[0],
    "contact": texts[1],
    "intent": texts[2].replace("求职意向：", ""),
    "overview": texts[4],
    "abilities": texts[6:11],
    "summary_label": texts[13].rstrip("："),
    "summary": texts[14],
    "duties": texts[15:19],
    "outcomes": [
        (texts[20], [texts[21]]),
        (texts[22], [texts[23]]),
        (texts[24], [texts[25]]),
        (texts[26], [texts[27]]),
        (texts[28], [texts[29]]),
        (texts[30], [texts[31]]),
        (texts[32], [texts[33], texts[34], texts[35], texts[36], texts[37]]),
        (texts[38], [texts[39]]),
        (texts[40], [texts[41]]),
        (texts[42], [texts[43]]),
        (texts[44], [texts[45]]),
    ],
    "ali": texts[46:52],
    "mid": texts[52:55],
    "state_grid": texts[55:58],
    "early": texts[58:60],
}


def drop_num(title: str) -> str:
    return title.split("、", 1)[1] if "、" in title else title


base = getSampleStyleSheet()


def ps(name, size, leading, color=INK, align=TA_LEFT, before=0, after=0):
    return ParagraphStyle(
        name,
        parent=base["Normal"],
        fontName=FONT,
        fontSize=size,
        leading=leading,
        textColor=color,
        alignment=align,
        spaceBefore=before,
        spaceAfter=after,
    )


S = {
    "name": ps("name", 28, 33, INK, TA_LEFT, after=2),
    "role": ps("role", 12.3, 16, CYAN, TA_LEFT, after=4),
    "contact": ps("contact", 8.5, 11, MUTED, TA_LEFT, after=6),
    "kicker": ps("kicker", 6.9, 9, CYAN, TA_LEFT, after=2),
    "h2": ps("h2", 13.5, 17, INK, TA_LEFT, before=5, after=5),
    "h3": ps("h3", 10.5, 13, INK, TA_LEFT, before=3, after=3),
    "body": ps("body", 8.25, 10.8, INK, TA_LEFT, after=2),
    "small": ps("small", 7.35, 9.7, MUTED, TA_LEFT, after=1.2),
    "tiny": ps("tiny", 6.8, 8.6, MUTED, TA_LEFT, after=0),
    "metric": ps("metric", 14, 17, INK, TA_CENTER),
    "metric_label": ps("metric_label", 6.7, 8.2, MUTED, TA_CENTER),
    "module": ps("module", 11.2, 13.5, INK, TA_LEFT, after=3),
    "outcome_title": ps("outcome_title", 8.7, 10.8, INK, TA_LEFT, after=1.5),
    "outcome_body": ps("outcome_body", 7.65, 9.7, INK, TA_LEFT, after=1.3),
    "date": ps("date", 8.0, 10, MUTED, TA_RIGHT),
}


def P(text: str, style="body"):
    return Paragraph(escape(text), S[style])


class SystemChain(Flowable):
    def __init__(self, width: float, height: float = 37 * mm):
        super().__init__()
        self.width = width
        self.height = height
        self.nodes = [
            ("复杂业务", "对象 / 流程 / 权限"),
            ("业务本体", "Ontology / 规则"),
            ("知识生产", "FMEA / 质检 / 复用"),
            ("AI复盘", "LLM / 多模态 / MCP"),
            ("Agent调度", "资源 / 任务 / 监测"),
            ("经营决策", "产能 / 策略 / 边界"),
        ]

    def wrap(self, availWidth, availHeight):
        return self.width, self.height

    def draw(self):
        c = self.canv
        c.saveState()
        c.setFillColor(PANEL)
        c.roundRect(0, 0, self.width, self.height, 5, fill=1, stroke=0)
        c.setStrokeColor(HAIR)
        c.roundRect(0, 0, self.width, self.height, 5, fill=0, stroke=1)
        c.setFont(FONT, 6.7)
        c.setFillColor(CYAN)
        c.drawString(7 * mm, self.height - 7 * mm, "CAPABILITY SYSTEM")
        left = 7 * mm
        top = self.height - 15 * mm
        gap = 2.2 * mm
        node_w = (self.width - 14 * mm - gap * 5) / 6
        node_h = 17 * mm
        colors_ = [GRAPHITE, BLUE, CYAN, TEAL, AMBER, PLUM]
        for i, (title, sub) in enumerate(self.nodes):
            x = left + i * (node_w + gap)
            c.setFillColor(colors_[i])
            c.roundRect(x, top - node_h, node_w, node_h, 3, fill=1, stroke=0)
            c.setFillColor(WHITE)
            c.setFont(FONT, 8.2)
            c.drawCentredString(x + node_w / 2, top - 6.4 * mm, title)
            c.setFont(FONT, 5.7)
            c.setFillColor(colors.HexColor("#ECF7FA"))
            c.drawCentredString(x + node_w / 2, top - 12.2 * mm, sub)
            if i < 5:
                x1 = x + node_w
                x2 = x + node_w + gap
                y = top - node_h / 2
                c.setStrokeColor(colors.HexColor("#8EA5B6"))
                c.setLineWidth(0.7)
                c.line(x1 + 1, y, x2 - 3, y)
                c.setFillColor(colors.HexColor("#8EA5B6"))
                c.circle(x2 - 2, y, 1.2, fill=1, stroke=0)
        c.restoreState()


def section(title: str, note: str | None = None):
    right = P(note or "", "tiny")
    table = Table(
        [[Paragraph(escape(title), S["h2"]), right]],
        colWidths=[118 * mm, 54 * mm],
        style=TableStyle(
            [
                ("LINEBELOW", (0, 0), (-1, -1), 0.6, HAIR),
                ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        ),
    )
    return [Spacer(1, 3), table, Spacer(1, 4)]


def intro_panel():
    return Table(
        [[P(data["overview"], "body")]],
        colWidths=[172 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7FAFB")),
                ("LINEBEFORE", (0, 0), (0, -1), 2.8, CYAN),
                ("BOX", (0, 0), (-1, -1), 0.4, HAIR),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        ),
    )


def metric(num, label, color=CYAN):
    return Table(
        [[Paragraph(escape(num), S["metric"])], [Paragraph(escape(label), S["metric_label"])]],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), WHITE),
                ("BOX", (0, 0), (-1, -1), 0.45, HAIR),
                ("LINEABOVE", (0, 0), (-1, 0), 2.0, color),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        ),
    )


def metric_strip():
    rows = [
        metric("980", "风险事件", BLUE),
        metric("5万+", "风险隐患节点", CYAN),
        metric("10万+", "观测及维修措施", TEAL),
        metric("7-10→2天", "复盘报告周期", AMBER),
        metric("95%", "调度建议采纳率", CYAN),
        metric("30%→1%", "工单阻塞率", PLUM),
    ]
    return Table(
        [rows],
        colWidths=[27.8 * mm] * 6,
        style=TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
            ]
        ),
    )


def ability_ladder():
    rows = []
    accents = [BLUE, CYAN, TEAL, AMBER, PLUM]
    for i, text in enumerate(data["abilities"]):
        title, body = text.split("：", 1)
        rows.append([
            Paragraph(f"{i+1:02d}", ParagraphStyle(f"idx{i}", fontName=FONT, fontSize=8.0, leading=10, textColor=WHITE, alignment=TA_CENTER)),
            Paragraph(escape(title), ParagraphStyle(f"at{i}", fontName=FONT, fontSize=9.1, leading=11.3, textColor=accents[i])),
            Paragraph(escape(body), ParagraphStyle(f"ab{i}", fontName=FONT, fontSize=7.55, leading=9.8, textColor=INK)),
        ])
    t = Table(rows, colWidths=[9 * mm, 36 * mm, 124 * mm])
    style = [
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOX", (0, 0), (-1, -1), 0.4, HAIR),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#E8EEF2")),
        ("BACKGROUND", (1, 0), (-1, -1), colors.HexColor("#FCFDFE")),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]
    for i, color in enumerate(accents):
        style.append(("BACKGROUND", (0, i), (0, i), color))
    t.setStyle(TableStyle(style))
    return t


def exp_header(company, role, date):
    return Table(
        [[Paragraph(escape(company), S["module"]), Paragraph(escape(role), S["small"]), Paragraph(escape(date), S["date"])]],
        colWidths=[74 * mm, 58 * mm, 37 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F3F7FA")),
                ("LINEBEFORE", (0, 0), (0, -1), 2.2, CYAN),
                ("BOX", (0, 0), (-1, -1), 0.45, HAIR),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        ),
    )


def bullet_list(items, font_style="body", width=166 * mm):
    rows = []
    for item in items:
        rows.append([
            Paragraph("•", ParagraphStyle("dot", fontName=FONT, fontSize=8, leading=10, textColor=CYAN)),
            P(item, font_style),
        ])
    t = Table(rows, colWidths=[5 * mm, width - 5 * mm])
    t.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    return t


def scope_matrix():
    labels = ["知识工程", "数字化产品", "Agent治理", "团队协同"]
    rows = []
    for i, duty in enumerate(data["duties"]):
        rows.append([Paragraph(labels[i], S["module"]), P(duty, "small")])
    t = Table(rows, colWidths=[26 * mm, 141 * mm])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EEF6F7")),
                ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#FBFCFD")),
                ("TEXTCOLOR", (0, 0), (0, -1), CYAN),
                ("BOX", (0, 0), (-1, -1), 0.4, HAIR),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#E7EEF2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return t


def outcome_item(num: int, title: str, bodies: list[str], accent):
    body = [Paragraph(escape(drop_num(title)), S["outcome_title"])]
    for b in bodies:
        body.append(Paragraph(escape(b), S["outcome_body"]))
    t = Table(
        [[Paragraph(f"{num:02d}", ParagraphStyle(f"n{num}", fontName=FONT, fontSize=8.2, leading=10, textColor=WHITE, alignment=TA_CENTER)), body]],
        colWidths=[9 * mm, 157 * mm],
    )
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), accent),
                ("BACKGROUND", (1, 0), (1, 0), WHITE),
                ("BOX", (0, 0), (-1, -1), 0.35, HAIR),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (0, 0), 3),
                ("RIGHTPADDING", (0, 0), (0, 0), 3),
                ("TOPPADDING", (0, 0), (0, 0), 6),
                ("LEFTPADDING", (1, 0), (1, 0), 7),
                ("RIGHTPADDING", (1, 0), (1, 0), 7),
                ("TOPPADDING", (1, 0), (1, 0), 5),
                ("BOTTOMPADDING", (1, 0), (1, 0), 5),
            ]
        )
    )
    return t


def campaign(title: str, premise: str, indices: list[int], accent):
    items = []
    for idx in indices:
        t, b = data["outcomes"][idx - 1]
        items.append(outcome_item(idx, t, b, accent))
        items.append(Spacer(1, 3))
    head = Table(
        [[Paragraph(escape(title), S["module"]), Paragraph(escape(premise), S["small"])]],
        colWidths=[46 * mm, 121 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), accent),
                ("TEXTCOLOR", (0, 0), (0, -1), WHITE),
                ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#F8FAFB")),
                ("BOX", (0, 0), (-1, -1), 0.4, HAIR),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        ),
    )
    return [KeepTogether([head, Spacer(1, 4)]), *items]


def ledger_block(head, items):
    return [exp_header(*head), bullet_list(items, "small"), Spacer(1, 5)]


def header_story():
    role = data["intent"].split(" | ")[0]
    return [
        Paragraph("AI PRODUCT SYSTEM ARCHITECTURE RESUME", S["kicker"]),
        Table(
            [[Paragraph(data["name"], S["name"]), Paragraph(escape(role), ParagraphStyle("role_r", fontName=FONT, fontSize=11.5, leading=15, textColor=CYAN, alignment=TA_RIGHT))]],
            colWidths=[72 * mm, 98 * mm],
            style=TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ]
            ),
        ),
        Paragraph(escape(data["contact"] + " | 期望薪资：40K"), S["contact"]),
    ]


def on_page(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(PAPER)
    canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    canvas.setStrokeColor(HAIR)
    canvas.setLineWidth(0.5)
    canvas.line(14 * mm, PAGE_H - 11 * mm, PAGE_W - 14 * mm, PAGE_H - 11 * mm)
    canvas.setFillColor(MUTED)
    canvas.setFont(FONT, 6.8)
    canvas.drawString(14 * mm, PAGE_H - 8 * mm, "曹宇迪 | B端AI应用产品负责人")
    canvas.drawRightString(PAGE_W - 14 * mm, PAGE_H - 8 * mm, f"PAGE {doc.page}")
    canvas.restoreState()


def build_pdf():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = BaseDocTemplate(
        str(PDF),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=15 * mm,
        bottomMargin=12 * mm,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=on_page)])

    story = []
    story.extend(header_story())
    story.append(SystemChain(172 * mm))
    story.append(Spacer(1, 7))
    story.extend(section("职业概览", "从复杂业务到 AI 可控执行"))
    story.append(intro_panel())
    story.append(Spacer(1, 7))
    story.append(metric_strip())
    story.extend(section("核心能力", "能力不是标签，是一条可验证的产品化链路"))
    story.append(ability_ladder())
    story.append(PageBreak())

    story.extend(section("主战场：企业级 AI 应用与数字化产品体系", "2021.02-至今"))
    story.append(exp_header(*headers[0]))
    story.append(Spacer(1, 4))
    story.append(
        Table(
            [[Paragraph(data["summary_label"], S["module"]), P(data["summary"], "body")]],
            colWidths=[25 * mm, 142 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8F5F5")),
                    ("BACKGROUND", (1, 0), (1, -1), WHITE),
                    ("BOX", (0, 0), (-1, -1), 0.45, HAIR),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ]
            ),
        )
    )
    story.extend(section("职责范围", "把职责压成四条产品系统边界"))
    story.append(scope_matrix())
    story.extend(section("三条证明链", "不是罗列成果，而是证明 AI 产品负责人能力如何形成"))
    story.extend(campaign("战役 01 · 知识工程", "从风险事件、设备实体、失效模式到结构化 FMEA 知识生产，解决 AI 应用冷启动和复杂方案表达问题。", [1, 2], BLUE))
    story.extend(campaign("战役 02 · 数字化业务闭环", "把配置、管理、执行、复盘和规模化客户落地连成可追踪的业务运行系统。", [3, 4, 6, 7], TEAL))
    story.append(Spacer(1, 6))
    story.extend(section("三条证明链（续）", "AI 复盘、Agent 调度与经营治理"))
    story.extend(campaign("战役 03 · AI/Agent 治理", "让 AI 从生成内容走向可监测、可验证、可辅助经营决策的企业级执行链路。", [5, 8, 9, 10, 11], CYAN))
    story.append(Spacer(1, 6))
    story.extend(section("平台能力与早期产品管理经历", "PaaS 中台、团队管理、商业模式产品化"))
    story.extend(ledger_block(headers[1], data["ali"]))
    story.extend(ledger_block(headers[2], data["mid"]))
    story.append(PageBreak())

    story.extend(section("早期经历与补充信息"))
    story.extend(ledger_block(headers[3], data["state_grid"]))
    story.extend(ledger_block(headers[4], data["early"]))
    story.extend(section("教育经历"))
    story.append(
        Table(
            [[P(headers[5][0], "module"), P(headers[5][1], "small"), Paragraph(escape(headers[5][2]), S["date"])]],
            colWidths=[83 * mm, 52 * mm, 34 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F5F8FA")),
                    ("BOX", (0, 0), (-1, -1), 0.45, HAIR),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            ),
        )
    )
    story.extend(section("认证"))
    story.append(Spacer(1, 12))
    story.extend(section("个人作品"))
    story.append(
        Table(
            [[
                Paragraph("PORTFOLIO<br/>LINK SLOT", ParagraphStyle("portfolio_mark", fontName=FONT, fontSize=16, leading=20, textColor=WHITE, alignment=TA_CENTER)),
                [
                    Paragraph("Portfolio / 个人作品链接：待补充", S["module"]),
                    Paragraph("这个区域预留给后续小网页链接。链接页可承接 AI 产品实践、Agent 工作流、设计工程方法论与可访问项目页面，让简历从静态履历延伸为可验证作品集。", S["small"]),
                    Spacer(1, 7),
                    Table(
                        [[
                            Paragraph("AI 产品实践", S["tiny"]),
                            Paragraph("Agent 工作流", S["tiny"]),
                            Paragraph("设计工程方法论", S["tiny"]),
                        ]],
                        colWidths=[31 * mm, 31 * mm, 36 * mm],
                        style=TableStyle(
                            [
                                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EEF6F7")),
                                ("BOX", (0, 0), (-1, -1), 0.35, HAIR),
                                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DCE8EE")),
                                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                ("TOPPADDING", (0, 0), (-1, -1), 5),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                            ]
                        ),
                    ),
                ],
            ]],
            colWidths=[48 * mm, 120 * mm],
            rowHeights=[48 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), GRAPHITE),
                    ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#F7FAFB")),
                    ("BOX", (0, 0), (-1, -1), 0.45, HAIR),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (0, 0), (0, -1), "CENTER"),
                    ("LEFTPADDING", (0, 0), (0, -1), 8),
                    ("RIGHTPADDING", (0, 0), (0, -1), 8),
                    ("TOPPADDING", (0, 0), (0, -1), 14),
                    ("LEFTPADDING", (1, 0), (1, -1), 12),
                    ("RIGHTPADDING", (1, 0), (1, -1), 12),
                    ("TOPPADDING", (1, 0), (1, -1), 11),
                    ("BOTTOMPADDING", (1, 0), (1, -1), 10),
                ]
            ),
        )
    )

    doc.build(story)


def build_html():
    def li(items):
        return "".join(f"<li>{escape(x)}</li>" for x in items)

    chain = "".join(f"<div><b>{escape(a)}</b><span>{escape(b)}</span></div>" for a, b in [
        ("复杂业务", "对象/流程/权限"),
        ("业务本体", "Ontology/规则"),
        ("知识生产", "FMEA/质检/复用"),
        ("AI复盘", "LLM/多模态/MCP"),
        ("Agent调度", "资源/任务/监测"),
        ("经营决策", "产能/策略/边界"),
    ])
    abilities = ""
    for i, text in enumerate(data["abilities"], 1):
        title, body = text.split("：", 1)
        abilities += f"<div class='ability'><span>{i:02d}</span><b>{escape(title)}</b><p>{escape(body)}</p></div>"

    def outcome_html(idx):
        title, bodies = data["outcomes"][idx - 1]
        return f"<article class='outcome'><i>{idx:02d}</i><div><h4>{escape(drop_num(title))}</h4>{''.join(f'<p>{escape(b)}</p>' for b in bodies)}</div></article>"

    html = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8" />
<title>曹宇迪简历-系统架构视觉版</title>
<style>
:root{{--ink:#151A1F;--graphite:#25313B;--muted:#64717D;--hair:#D9E1E7;--paper:#FAFBFC;--panel:#F2F6F8;--teal:#0F766E;--cyan:#0E7490;--blue:#1D4E89;--amber:#A16207;--plum:#6D4C7D;}}
*{{box-sizing:border-box}} body{{margin:0;background:#edf2f5;color:var(--ink);font-family:-apple-system,BlinkMacSystemFont,"PingFang SC","Microsoft YaHei",sans-serif;}}
.page{{width:210mm;min-height:297mm;margin:18px auto;background:var(--paper);box-shadow:0 14px 40px rgba(21,26,31,.12);padding:17mm 18mm;}}
.kicker{{font-size:10px;color:var(--cyan);letter-spacing:.12em}} header{{display:flex;justify-content:space-between;align-items:flex-end;border-bottom:1px solid var(--hair);padding-bottom:10px}}
h1{{font-size:38px;margin:4px 0;color:var(--ink)}} .role{{font-size:16px;color:var(--cyan);text-align:right}} .contact{{font-size:12px;color:var(--muted);margin-top:7px}}
.chain{{display:grid;grid-template-columns:repeat(6,1fr);gap:7px;background:var(--panel);border:1px solid var(--hair);padding:13px;margin:16px 0}}
.chain div{{color:white;padding:12px 8px;min-height:58px;background:var(--blue)}} .chain div:nth-child(2){{background:var(--cyan)}} .chain div:nth-child(3){{background:var(--teal)}} .chain div:nth-child(4){{background:var(--amber)}} .chain div:nth-child(5){{background:var(--plum)}} .chain span{{display:block;font-size:10px;opacity:.9;margin-top:6px}}
h2{{font-size:20px;border-bottom:1px solid var(--hair);padding-bottom:7px;margin:20px 0 12px}} h3{{font-size:16px;margin:14px 0 8px}}
.overview,.summary,.portfolio{{border-left:4px solid var(--cyan);background:white;border:1px solid var(--hair);padding:14px 16px;line-height:1.75;font-size:13px}}
.metrics{{display:grid;grid-template-columns:repeat(6,1fr);gap:7px;margin:14px 0}} .metric{{background:white;border:1px solid var(--hair);border-top:3px solid var(--cyan);text-align:center;padding:10px 2px}} .metric b{{font-size:19px}} .metric span{{display:block;font-size:10px;color:var(--muted);margin-top:5px}}
.ability{{display:grid;grid-template-columns:34px 130px 1fr;border:1px solid var(--hair);border-bottom:0;background:white}} .ability:last-child{{border-bottom:1px solid var(--hair)}} .ability span{{background:var(--cyan);color:white;text-align:center;padding-top:14px}} .ability b{{color:var(--cyan);padding:12px}} .ability p{{font-size:12px;line-height:1.65;margin:0;padding:12px}}
.exp{{display:grid;grid-template-columns:1.3fr 1fr auto;gap:12px;background:#F3F7FA;border-left:4px solid var(--cyan);border:1px solid var(--hair);padding:10px 12px;color:var(--muted)}} .exp b{{color:var(--ink)}}
.scope{{display:grid;grid-template-columns:120px 1fr;border:1px solid var(--hair)}} .scope b{{background:#E8F5F5;padding:12px;color:var(--cyan)}} .scope p{{margin:0;padding:12px;border-bottom:1px solid var(--hair);font-size:12px;line-height:1.65}}
.campaign{{margin-top:12px;border:1px solid var(--hair)}} .campaign-head{{display:grid;grid-template-columns:190px 1fr}} .campaign-head b{{background:var(--cyan);color:white;padding:13px}} .campaign-head span{{background:white;padding:13px;color:var(--muted);line-height:1.6;font-size:12px}}
.outcome{{display:grid;grid-template-columns:36px 1fr;border-top:1px solid var(--hair);background:white}} .outcome i{{background:var(--graphite);color:white;text-align:center;padding-top:12px;font-style:normal}} .outcome div{{padding:10px 12px}} .outcome h4{{margin:0 0 5px;font-size:13px}} .outcome p{{font-size:11.5px;line-height:1.68;margin:0 0 4px}}
ul{{margin:8px 0 14px 20px;padding:0}} li{{font-size:12px;line-height:1.7;margin-bottom:4px}}
@media print{{body{{background:white}}.page{{margin:0;box-shadow:none;page-break-after:always}}}}
</style></head><body>
<section class="page"><div class="kicker">AI PRODUCT SYSTEM ARCHITECTURE RESUME</div><header><div><h1>{escape(data["name"])}</h1><div class="contact">{escape(data["contact"])} | 期望薪资：40K</div></div><div class="role">{escape(data["intent"].split(" | ")[0])}</div></header><div class="chain">{chain}</div><h2>职业概览</h2><div class="overview">{escape(data["overview"])}</div><div class="metrics"><div class="metric"><b>980</b><span>风险事件</span></div><div class="metric"><b>5万+</b><span>风险隐患节点</span></div><div class="metric"><b>10万+</b><span>观测及维修措施</span></div><div class="metric"><b>7-10→2天</b><span>复盘报告周期</span></div><div class="metric"><b>95%</b><span>调度建议采纳率</span></div><div class="metric"><b>30%→1%</b><span>工单阻塞率</span></div></div><h2>核心能力</h2>{abilities}</section>
<section class="page"><h2>主战场：企业级 AI 应用与数字化产品体系</h2><div class="exp"><b>{escape(headers[0][0])}</b><span>{escape(headers[0][1])}</span><span>{escape(headers[0][2])}</span></div><h3>{escape(data["summary_label"])}</h3><div class="summary">{escape(data["summary"])}</div><h2>职责范围</h2><div class="scope">{"".join(f"<b>{escape(t)}</b><p>{escape(x)}</p>" for t,x in zip(["知识工程","数字化产品","Agent治理","团队协同"],data["duties"]))}</div><h2>三条证明链</h2><div class="campaign"><div class="campaign-head"><b>战役 01 · 知识工程</b><span>从风险事件、设备实体、失效模式到结构化 FMEA 知识生产，解决 AI 应用冷启动和复杂方案表达问题。</span></div>{outcome_html(1)}{outcome_html(2)}</div><div class="campaign"><div class="campaign-head"><b>战役 02 · 数字化业务闭环</b><span>把配置、管理、执行、复盘和规模化客户落地连成可追踪的业务运行系统。</span></div>{outcome_html(3)}{outcome_html(4)}{outcome_html(6)}{outcome_html(7)}</div></section>
<section class="page"><h2>三条证明链（续）</h2><div class="campaign"><div class="campaign-head"><b>战役 03 · AI/Agent 治理</b><span>让 AI 从生成内容走向可监测、可验证、可辅助经营决策的企业级执行链路。</span></div>{outcome_html(5)}{outcome_html(8)}{outcome_html(9)}{outcome_html(10)}{outcome_html(11)}</div><h2>平台能力与早期产品管理经历</h2><div class="exp"><b>{escape(headers[1][0])}</b><span>{escape(headers[1][1])}</span><span>{escape(headers[1][2])}</span></div><ul>{li(data["ali"])}</ul><div class="exp"><b>{escape(headers[2][0])}</b><span>{escape(headers[2][1])}</span><span>{escape(headers[2][2])}</span></div><ul>{li(data["mid"])}</ul></section>
<section class="page"><h2>早期经历与补充信息</h2><div class="exp"><b>{escape(headers[3][0])}</b><span>{escape(headers[3][1])}</span><span>{escape(headers[3][2])}</span></div><ul>{li(data["state_grid"])}</ul><div class="exp"><b>{escape(headers[4][0])}</b><span>{escape(headers[4][1])}</span><span>{escape(headers[4][2])}</span></div><ul>{li(data["early"])}</ul><h2>教育经历</h2><div class="exp"><b>{escape(headers[5][0])}</b><span>{escape(headers[5][1])}</span><span>{escape(headers[5][2])}</span></div><h2>认证</h2><h2>个人作品</h2><div class="portfolio" style="display:grid;grid-template-columns:180px 1fr;min-height:180px;padding:0"><div style="background:var(--graphite);color:white;padding:40px 24px;font-size:24px;line-height:1.25">PORTFOLIO<br/>LINK SLOT</div><div style="padding:24px"><b>Portfolio / 个人作品链接：待补充</b><br/><br/>这个区域预留给后续小网页链接。链接页可承接 AI 产品实践、Agent 工作流、设计工程方法论与可访问项目页面。<div style="display:flex;gap:8px;margin-top:18px"><span style="border:1px solid var(--hair);background:#EEF6F7;padding:8px 12px">AI 产品实践</span><span style="border:1px solid var(--hair);background:#EEF6F7;padding:8px 12px">Agent 工作流</span><span style="border:1px solid var(--hair);background:#EEF6F7;padding:8px 12px">设计工程方法论</span></div></div></div></section>
</body></html>"""
    HTML.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build_html()
    build_pdf()
    print(HTML)
    print(PDF)
