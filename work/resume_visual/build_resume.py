from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path("/Users/caoyudi/Documents/Codex/2026-07-03/ni-xai")
SRC = Path("/Users/caoyudi/Downloads/曹宇迪简历.docx")
OUT_DIR = ROOT / "outputs"
PDF_OUT = OUT_DIR / "曹宇迪简历-视觉设计版.pdf"
HTML_OUT = OUT_DIR / "曹宇迪简历-视觉设计版.html"

FONT = "STHeiti"
FONT_PATH = "/System/Library/Fonts/STHeiti Medium.ttc"
pdfmetrics.registerFont(TTFont(FONT, FONT_PATH))

PAGE_W, PAGE_H = A4

NAVY = colors.HexColor("#102A43")
BLUE = colors.HexColor("#1D5F90")
BLUE_2 = colors.HexColor("#2F80B7")
CYAN = colors.HexColor("#58A6C7")
INK = colors.HexColor("#1F2933")
MUTED = colors.HexColor("#607080")
LIGHT = colors.HexColor("#F4F8FB")
LIGHT_2 = colors.HexColor("#EAF2F8")
BLUE_SOFT = colors.HexColor("#EAF2F8")
LINE = colors.HexColor("#D8E4EE")
WHITE = colors.white
SOFT_GREEN = colors.HexColor("#E9F7F2")
GREEN = colors.HexColor("#2F7D62")
SOFT_GOLD = colors.HexColor("#FFF5DF")
GOLD = colors.HexColor("#9A6A00")


def clean(text: str) -> str:
    text = " ".join(text.split())
    text = text.replace("负责职责说明：", "负责。")
    text = text.replace(" 半天以内", "半天以内")
    return text


def load_resume():
    doc = DocxDocument(str(SRC))
    texts = [clean(p.text.strip()) for p in doc.paragraphs]
    tables = []
    for table in doc.tables:
        tables.append([clean(cell.text) for cell in table.rows[0].cells])
    return texts, tables


texts, headers = load_resume()

profile = {
    "name": texts[0],
    "contact": texts[1],
    "intent": texts[2].replace("求职意向：", ""),
    "overview": texts[4],
    "abilities": texts[6:11],
    "work_summary_label": texts[13].rstrip("："),
    "work_summary": texts[14],
    "duties": texts[15:19],
    "outcomes": [
        (texts[20], [texts[21]]),
        (texts[22], [texts[23]]),
        (texts[24], [texts[25]]),
        (texts[26], [texts[27]]),
        (texts[28], [texts[29]]),
        (texts[30], [texts[31]]),
        (texts[32], [texts[33], texts[34], texts[35], texts[36], texts[37]]),
        (texts[38], [texts[39]]),
        (texts[40], [texts[41]]),
        (texts[42], [texts[43]]),
        (texts[44], [texts[45]]),
    ],
    "ali": texts[46:52],
    "mid": texts[52:55],
    "state_grid": texts[55:58],
    "early": texts[58:60],
}


def split_title(text: str) -> str:
    return text.split("、", 1)[1] if "、" in text else text


def stylesheet():
    base = getSampleStyleSheet()
    styles = {
        "h1": ParagraphStyle(
            "h1",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=25,
            leading=30,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=3,
        ),
        "role": ParagraphStyle(
            "role",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=11.5,
            leading=15,
            textColor=BLUE,
            alignment=TA_CENTER,
            spaceAfter=4,
        ),
        "contact": ParagraphStyle(
            "contact",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=8.8,
            leading=12,
            textColor=MUTED,
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "section": ParagraphStyle(
            "section",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=13,
            leading=16,
            textColor=NAVY,
            spaceBefore=8,
            spaceAfter=5,
        ),
        "subsection": ParagraphStyle(
            "subsection",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=10.5,
            leading=13,
            textColor=BLUE,
            spaceBefore=4,
            spaceAfter=3,
        ),
        "body": ParagraphStyle(
            "body",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=8.55,
            leading=11.4,
            textColor=INK,
            spaceAfter=2.2,
        ),
        "small": ParagraphStyle(
            "small",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=7.8,
            leading=10.4,
            textColor=MUTED,
            spaceAfter=1.5,
        ),
        "card_title": ParagraphStyle(
            "card_title",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=9.1,
            leading=11.2,
            textColor=NAVY,
            spaceAfter=2,
        ),
        "card_body": ParagraphStyle(
            "card_body",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=8.0,
            leading=10.6,
            textColor=INK,
            spaceAfter=1.5,
        ),
        "metric": ParagraphStyle(
            "metric",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=15,
            leading=18,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=1,
        ),
        "metric_label": ParagraphStyle(
            "metric_label",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=6.8,
            leading=8,
            textColor=MUTED,
            alignment=TA_CENTER,
        ),
        "white_big": ParagraphStyle(
            "white_big",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=22,
            leading=26,
            textColor=WHITE,
            spaceAfter=3,
        ),
        "white": ParagraphStyle(
            "white",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=8.1,
            leading=11,
            textColor=colors.HexColor("#DDEAF3"),
            spaceAfter=2,
        ),
        "white_label": ParagraphStyle(
            "white_label",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=7.5,
            leading=10,
            textColor=colors.HexColor("#9ED0E8"),
            spaceBefore=7,
            spaceAfter=2,
        ),
    }
    return styles


S = stylesheet()


def P(text: str, style: str = "body"):
    return Paragraph(escape(text), S[style])


def section(title: str, width: float = 172 * mm):
    return [
        Spacer(1, 3),
        Table(
            [[P(title, "section")]],
            colWidths=[width],
            style=TableStyle(
                [
                    ("LINEBELOW", (0, 0), (-1, -1), 0.8, LINE),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                    ("TOPPADDING", (0, 0), (-1, -1), 1),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ]
            ),
        ),
    ]


def pill(text: str, bg=LIGHT_2, fg=BLUE):
    return Table(
        [[Paragraph(escape(text), ParagraphStyle("pill", fontName=FONT, fontSize=7.3, leading=9, textColor=fg, alignment=TA_CENTER))]],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), bg),
                ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#C9DCEB")),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        ),
    )


def overview_box(width: float = 172 * mm):
    return Table(
        [[P(profile["overview"], "body")]],
        colWidths=[width],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), LIGHT),
                ("BOX", (0, 0), (-1, -1), 0.5, LINE),
                ("LINEBEFORE", (0, 0), (0, -1), 3, BLUE),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        ),
    )


def metric_card(num: str, label: str):
    return Table(
        [[Paragraph(num, S["metric"])], [Paragraph(label, S["metric_label"])]],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), WHITE),
                ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        ),
    )


def ability_cell(text: str):
    title, body = text.split("：", 1)
    return [
        Paragraph(escape(title), ParagraphStyle("ability_title", fontName=FONT, fontSize=8.9, leading=11, textColor=BLUE)),
        Paragraph(escape(body), ParagraphStyle("ability_body", fontName=FONT, fontSize=7.65, leading=10.1, textColor=INK)),
    ]


def ability_matrix(col_width: float = 84 * mm):
    rows = [
        [ability_cell(profile["abilities"][0]), ability_cell(profile["abilities"][1])],
        [ability_cell(profile["abilities"][2]), ability_cell(profile["abilities"][3])],
        [ability_cell(profile["abilities"][4]), ""],
    ]
    t = Table(rows, colWidths=[col_width, col_width])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FBFCFD")),
                ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#E3ECF3")),
                ("SPAN", (0, 2), (1, 2)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return t


def exp_header(company: str, role: str, dates: str):
    return Table(
        [[P(company, "card_title"), Paragraph(escape(role), S["small"]), Paragraph(escape(dates), ParagraphStyle("date", fontName=FONT, fontSize=8.4, leading=11, textColor=MUTED, alignment=TA_RIGHT))]],
        colWidths=[72 * mm, 62 * mm, 36 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), BLUE_SOFT),
                ("LINEBEFORE", (0, 0), (0, -1), 2.5, BLUE),
                ("BOX", (0, 0), (-1, -1), 0.4, LINE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        ),
    )


def bullet_table(items, size_style="body"):
    rows = []
    for item in items:
        rows.append([Paragraph("•", ParagraphStyle("dot", fontName=FONT, fontSize=8.3, leading=10, textColor=BLUE)), P(item, size_style)])
    t = Table(rows, colWidths=[5 * mm, 164 * mm])
    t.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 1.2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1.2),
            ]
        )
    )
    return t


def outcome_card(i: int, title: str, bodies: list[str], accent=BLUE):
    body_flow = [Paragraph(escape(split_title(title)), S["card_title"])]
    for body in bodies:
        body_flow.append(Paragraph(escape(body), S["card_body"]))
    number = Paragraph(str(i).zfill(2), ParagraphStyle("num", fontName=FONT, fontSize=9, leading=11, textColor=WHITE, alignment=TA_CENTER))
    t = Table([[number, body_flow]], colWidths=[10 * mm, 158 * mm])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), accent),
                ("BACKGROUND", (1, 0), (1, 0), WHITE),
                ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                ("LINEBEFORE", (1, 0), (1, -1), 0.8, colors.HexColor("#E3ECF3")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (0, 0), 3),
                ("RIGHTPADDING", (0, 0), (0, 0), 3),
                ("TOPPADDING", (0, 0), (0, 0), 6),
                ("LEFTPADDING", (1, 0), (1, 0), 8),
                ("RIGHTPADDING", (1, 0), (1, 0), 8),
                ("TOPPADDING", (1, 0), (1, 0), 6),
                ("BOTTOMPADDING", (1, 0), (1, 0), 6),
            ]
        )
    )
    return [t, Spacer(1, 3.2)]


def sidebar():
    chips = [
        "复杂业务本体建模",
        "FMEA 知识工程",
        "Agent 调度治理",
        "Human-in-the-loop",
        "PaaS / SaaS / B端AI",
    ]
    content = [
        Paragraph("PROFILE", S["white_label"]),
        Paragraph(profile["name"], S["white_big"]),
        Paragraph("B端AI应用产品负责人<br/>企业级AI产品负责人<br/>Agent产品负责人", S["white"]),
        Spacer(1, 8),
        Paragraph("CONTACT", S["white_label"]),
        Paragraph(profile["contact"].replace(" | ", "<br/>"), S["white"]),
        Paragraph("期望薪资：40K", S["white"]),
        Spacer(1, 8),
        Paragraph("SIGNALS", S["white_label"]),
        Paragraph("12年产品经验<br/>5年产品与团队管理<br/>95%调度建议采纳率<br/>30%→1%工单阻塞率<br/>7-10天→2天报告周期", S["white"]),
        Spacer(1, 8),
        Paragraph("KEYWORDS", S["white_label"]),
    ]
    for chip in chips:
        content.append(Paragraph("• " + escape(chip), S["white"]))
    return content


def first_page():
    main_width = 124 * mm
    main = [
        Paragraph(profile["name"], S["h1"]),
        Paragraph(profile["intent"].split(" | ")[0], S["role"]),
        Paragraph(profile["contact"] + " | 期望薪资：40K", S["contact"]),
        overview_box(main_width),
        Spacer(1, 6),
        Table(
            [[metric_card("12年", "产品规划 / 设计 / 落地"), metric_card("5年", "产品管理 / 团队管理"), metric_card("95%", "Agent调度建议采纳率"), metric_card("30%→1%", "工单阻塞率下降")]],
            colWidths=[30.5 * mm] * 4,
            style=TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 2), ("RIGHTPADDING", (0, 0), (-1, -1), 2)]),
        ),
        Spacer(1, 7),
        *section("核心能力", main_width),
        ability_matrix(62 * mm),
    ]
    t = Table([[sidebar(), main]], colWidths=[45 * mm, 134 * mm])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), NAVY),
                ("BACKGROUND", (1, 0), (1, 0), WHITE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (0, 0), 8),
                ("RIGHTPADDING", (0, 0), (0, 0), 8),
                ("TOPPADDING", (0, 0), (0, 0), 11),
                ("BOTTOMPADDING", (0, 0), (0, 0), 11),
                ("LEFTPADDING", (1, 0), (1, 0), 9),
                ("RIGHTPADDING", (1, 0), (1, 0), 0),
                ("TOPPADDING", (1, 0), (1, 0), 0),
                ("BOTTOMPADDING", (1, 0), (1, 0), 0),
            ]
        )
    )
    return [t]


def page_canvas(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(colors.HexColor("#F7FAFC"))
    canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    canvas.setFillColor(NAVY)
    canvas.rect(0, PAGE_H - 11 * mm, PAGE_W, 11 * mm, fill=1, stroke=0)
    canvas.setFillColor(colors.HexColor("#9ED0E8"))
    canvas.setFont(FONT, 7.2)
    canvas.drawString(15 * mm, PAGE_H - 7 * mm, "曹宇迪 | B端AI应用产品负责人")
    canvas.drawRightString(PAGE_W - 15 * mm, PAGE_H - 7 * mm, f"PAGE {doc.page}")
    canvas.restoreState()


def build_pdf():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = BaseDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=12 * mm,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="all", frames=[frame], onPage=page_canvas)])

    story = []
    story.extend(first_page())
    story.append(PageBreak())

    story.extend(section("主战场：企业级 AI 应用与数字化产品体系"))
    story.append(exp_header(*headers[0]))
    story.append(Spacer(1, 4))
    story.append(Table(
        [[P(profile["work_summary_label"], "subsection"), P(profile["work_summary"], "body")]],
        colWidths=[20 * mm, 148 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), LIGHT),
                ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                ("LINEBEFORE", (0, 0), (0, -1), 2.4, BLUE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        ),
    ))
    story.append(Spacer(1, 6))
    story.append(P("职责说明", "subsection"))
    story.append(bullet_table(profile["duties"], "body"))
    story.append(Spacer(1, 7))
    story.append(P("关键成果", "subsection"))

    accents = [BLUE, BLUE, BLUE, CYAN, GREEN, GREEN, GOLD, BLUE_2, BLUE_2, GOLD, NAVY]
    for i, (title, bodies) in enumerate(profile["outcomes"], 1):
        story.extend(outcome_card(i, title, bodies, accents[i - 1]))
        if i == 5:
            story.append(PageBreak())
            story.extend(section("关键成果（续）"))

    story.append(Spacer(1, 8))
    story.extend(section("平台能力与早期产品管理经历"))
    story.append(exp_header(*headers[1]))
    story.append(bullet_table(profile["ali"], "body"))
    story.append(Spacer(1, 5))
    story.append(exp_header(*headers[2]))
    story.append(bullet_table(profile["mid"], "body"))
    story.append(Spacer(1, 5))
    story.append(exp_header(*headers[3]))
    story.append(bullet_table(profile["state_grid"], "body"))
    story.append(Spacer(1, 5))
    story.append(exp_header(*headers[4]))
    story.append(bullet_table(profile["early"], "body"))
    story.extend(section("教育经历"))
    story.append(
        Table(
            [[P(headers[5][0], "card_title"), P(headers[5][1], "small"), Paragraph(escape(headers[5][2]), ParagraphStyle("edu_date", fontName=FONT, fontSize=8.2, leading=11, textColor=MUTED, alignment=TA_RIGHT))]],
            colWidths=[82 * mm, 52 * mm, 34 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), LIGHT),
                    ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ]
            ),
        )
    )
    story.extend(section("认证"))
    story.append(Spacer(1, 10))
    story.extend(section("个人作品"))
    story.append(
        Table(
            [[P("Portfolio / 个人作品链接：待补充", "card_title"), P("预留为在线作品集入口，可承接 AI 产品实践、Agent 工作流、设计工程方法论与可访问项目页面。", "small")]],
            colWidths=[68 * mm, 100 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), LIGHT),
                    ("BOX", (0, 0), (-1, -1), 0.45, LINE),
                    ("LINEBEFORE", (0, 0), (0, -1), 2.5, BLUE),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ]
            ),
        )
    )

    doc.build(story)


def html_section(title: str) -> str:
    return f'<h2>{escape(title)}</h2>'


def html_outcome(i: int, title: str, bodies: list[str]) -> str:
    body = "".join(f"<p>{escape(x)}</p>" for x in bodies)
    return f'<article class="outcome"><div class="num">{i:02d}</div><div><h3>{escape(split_title(title))}</h3>{body}</div></article>'


def build_html():
    def ability_html(text: str) -> str:
        title, body = text.split("：", 1)
        return f"<div class='ability'><b>{escape(title)}</b><p>{escape(body)}</p></div>"

    outcomes = "\n".join(html_outcome(i, t, b) for i, (t, b) in enumerate(profile["outcomes"], 1))
    ali = "".join(f"<li>{escape(x)}</li>" for x in profile["ali"])
    mid = "".join(f"<li>{escape(x)}</li>" for x in profile["mid"])
    state = "".join(f"<li>{escape(x)}</li>" for x in profile["state_grid"])
    early = "".join(f"<li>{escape(x)}</li>" for x in profile["early"])
    duties = "".join(f"<li>{escape(x)}</li>" for x in profile["duties"])
    abilities = "\n".join(ability_html(x) for x in profile["abilities"])
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>曹宇迪简历-视觉设计版</title>
  <style>
    :root {{
      --navy:#102A43; --blue:#1D5F90; --cyan:#58A6C7; --ink:#1F2933; --muted:#607080;
      --line:#D8E4EE; --light:#F4F8FB; --paper:#FFFFFF;
    }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; background:#edf3f8; color:var(--ink); font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Microsoft YaHei", sans-serif; }}
    .page {{ width:210mm; min-height:297mm; margin:18px auto; background:var(--paper); box-shadow:0 12px 40px rgba(16,42,67,.13); padding:16mm 17mm; }}
    .hero {{ display:grid; grid-template-columns:48mm 1fr; gap:12mm; }}
    .side {{ background:var(--navy); color:#dcebf5; padding:11mm 8mm; min-height:250mm; }}
    .side h1 {{ color:white; font-size:26px; margin:0 0 8px; }}
    .side .role {{ color:#9ED0E8; font-size:12px; line-height:1.7; margin-bottom:18px; }}
    .side h4 {{ color:#9ED0E8; font-size:10px; letter-spacing:.12em; margin:18px 0 6px; }}
    .side p {{ font-size:11px; line-height:1.75; margin:0; }}
    .main h1 {{ color:var(--navy); text-align:center; font-size:32px; margin:0; }}
    .main .intent {{ color:var(--blue); text-align:center; font-size:15px; margin:5px 0 4px; }}
    .contact {{ color:var(--muted); text-align:center; font-size:12px; margin-bottom:12px; }}
    .overview {{ border-left:5px solid var(--blue); background:var(--light); padding:14px 16px; font-size:13px; line-height:1.8; }}
    .metrics {{ display:grid; grid-template-columns:repeat(4,1fr); gap:8px; margin:12px 0 16px; }}
    .metric {{ border:1px solid var(--line); text-align:center; padding:10px 4px; }}
    .metric b {{ color:var(--navy); font-size:20px; display:block; }}
    .metric span {{ color:var(--muted); font-size:10px; }}
    h2 {{ color:var(--navy); border-bottom:1px solid var(--line); padding-bottom:6px; margin:18px 0 10px; font-size:18px; }}
    h3 {{ margin:0 0 5px; color:var(--navy); font-size:14px; }}
    .abilities {{ display:grid; grid-template-columns:1fr 1fr; border:1px solid var(--line); }}
    .ability {{ padding:10px; border-right:1px solid var(--line); border-bottom:1px solid var(--line); background:#fbfcfd; }}
    .ability:nth-child(2n) {{ border-right:0; }}
    .ability:last-child {{ grid-column:1 / span 2; border-bottom:0; }}
    .ability b {{ color:var(--blue); font-size:13px; }}
    .ability p {{ margin:5px 0 0; font-size:11.3px; line-height:1.7; }}
    .exp-head {{ display:grid; grid-template-columns: 1.3fr 1fr auto; gap:10px; background:#EAF2F8; border-left:5px solid var(--blue); padding:9px 12px; margin-top:10px; align-items:center; }}
    .exp-head b {{ color:var(--navy); }} .exp-head span {{ color:var(--muted); font-size:12px; }}
    ul {{ margin:8px 0 12px 18px; padding:0; }} li {{ font-size:12px; line-height:1.72; margin-bottom:4px; }}
    .summary {{ background:var(--light); border-left:5px solid var(--blue); padding:12px 14px; font-size:12.5px; line-height:1.8; }}
    .outcome {{ display:grid; grid-template-columns:34px 1fr; border:1px solid var(--line); margin:8px 0; background:white; }}
    .outcome .num {{ background:var(--blue); color:white; text-align:center; padding-top:10px; font-size:13px; }}
    .outcome div:nth-child(2) {{ padding:10px 12px; }}
    .outcome p {{ margin:0 0 5px; font-size:11.5px; line-height:1.68; }}
    .portfolio {{ border-left:5px solid var(--blue); background:var(--light); padding:12px 14px; }}
    @media print {{ body {{ background:white; }} .page {{ margin:0; box-shadow:none; page-break-after:always; }} }}
  </style>
</head>
<body>
  <section class="page hero">
    <aside class="side">
      <h4>PROFILE</h4>
      <h1>{escape(profile["name"])}</h1>
      <div class="role">B端AI应用产品负责人<br/>企业级AI产品负责人<br/>Agent产品负责人</div>
      <h4>CONTACT</h4><p>{escape(profile["contact"]).replace(" | ", "<br/>")}<br/>期望薪资：40K</p>
      <h4>SIGNALS</h4><p>12年产品经验<br/>5年产品与团队管理<br/>95%调度建议采纳率<br/>30%→1%工单阻塞率<br/>7-10天→2天报告周期</p>
      <h4>KEYWORDS</h4><p>复杂业务本体建模<br/>FMEA 知识工程<br/>Agent 调度治理<br/>Human-in-the-loop<br/>PaaS / SaaS / B端AI</p>
    </aside>
    <main class="main">
      <h1>{escape(profile["name"])}</h1>
      <div class="intent">{escape(profile["intent"].split(" | ")[0])}</div>
      <div class="contact">{escape(profile["contact"])} | 期望薪资：40K</div>
      <div class="overview">{escape(profile["overview"])}</div>
      <div class="metrics"><div class="metric"><b>12年</b><span>产品规划 / 设计 / 落地</span></div><div class="metric"><b>5年</b><span>产品管理 / 团队管理</span></div><div class="metric"><b>95%</b><span>Agent调度建议采纳率</span></div><div class="metric"><b>30%→1%</b><span>工单阻塞率下降</span></div></div>
      {html_section("核心能力")}<div class="abilities">{abilities}</div>
    </main>
  </section>
  <section class="page">
    {html_section("主战场：企业级 AI 应用与数字化产品体系")}
    <div class="exp-head"><b>{escape(headers[0][0])}</b><span>{escape(headers[0][1])}</span><span>{escape(headers[0][2])}</span></div>
    <h3>{escape(profile["work_summary_label"])}</h3><div class="summary">{escape(profile["work_summary"])}</div>
    <h3>职责说明</h3><ul>{duties}</ul>
    <h3>关键成果</h3>{outcomes}
  </section>
  <section class="page">
    {html_section("平台能力与早期产品管理经历")}
    <div class="exp-head"><b>{escape(headers[1][0])}</b><span>{escape(headers[1][1])}</span><span>{escape(headers[1][2])}</span></div><ul>{ali}</ul>
    <div class="exp-head"><b>{escape(headers[2][0])}</b><span>{escape(headers[2][1])}</span><span>{escape(headers[2][2])}</span></div><ul>{mid}</ul>
    <div class="exp-head"><b>{escape(headers[3][0])}</b><span>{escape(headers[3][1])}</span><span>{escape(headers[3][2])}</span></div><ul>{state}</ul>
    <div class="exp-head"><b>{escape(headers[4][0])}</b><span>{escape(headers[4][1])}</span><span>{escape(headers[4][2])}</span></div><ul>{early}</ul>
    {html_section("教育经历")}<div class="exp-head"><b>{escape(headers[5][0])}</b><span>{escape(headers[5][1])}</span><span>{escape(headers[5][2])}</span></div>
    {html_section("认证")}<div style="height:18px"></div>
    {html_section("个人作品")}<div class="portfolio"><b>Portfolio / 个人作品链接：待补充</b><br/><span>预留为在线作品集入口，可承接 AI 产品实践、Agent 工作流、设计工程方法论与可访问项目页面。</span></div>
  </section>
</body>
</html>"""
    HTML_OUT.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_html()
    build_pdf()
    print(HTML_OUT)
    print(PDF_OUT)
